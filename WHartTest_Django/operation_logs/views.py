from rest_framework import viewsets, permissions, filters, status
from django_filters.rest_framework import DjangoFilterBackend
import django_filters
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.decorators import action
from rest_framework.parsers import MultiPartParser, JSONParser
from django.http import HttpResponse
import io
import os
import urllib.parse
import logging
import re as re_module

from .models import OperationLog, OperationLogSetting, AnonymizationRule, AnonymizedDocument, AnonymizationTemplate
from .serializers import (
    OperationLogSerializer, OperationLogSettingSerializer,
    AnonymizationRuleSerializer, AnonymizedDocumentSerializer,
    AnonymizationTemplateSerializer,
)
from .tasks import cleanup_operation_logs

logger = logging.getLogger(__name__)

class OperationLogFilter(django_filters.FilterSet):
    username = django_filters.CharFilter(lookup_expr='icontains')
    module = django_filters.CharFilter(lookup_expr='icontains')
    action = django_filters.CharFilter(lookup_expr='icontains')
    method = django_filters.CharFilter(lookup_expr='iexact')
    response_code = django_filters.NumberFilter()
    start_time = django_filters.DateTimeFilter(field_name='created_at', lookup_expr='gte')
    end_time = django_filters.DateTimeFilter(field_name='created_at', lookup_expr='lte')

    class Meta:
        model = OperationLog
        fields = ['username', 'module', 'action', 'method', 'response_code']

from wharttest_django.pagination import StandardPagination


class OperationLogReadPermission(permissions.BasePermission):
    def has_permission(self, request, view):
        if not request.user or request.user.is_anonymous:
            return False
        return (
            request.user.is_staff or
            request.user.has_perm('operation_logs.view_operationlog') or
            request.user.has_perm('accounts.view_operationlog')
        )


class OperationLogCleanupPermission(permissions.BasePermission):
    def has_permission(self, request, view):
        if not request.user or request.user.is_anonymous:
            return False
        return (
            request.user.is_staff or
            request.user.has_perm('operation_logs.delete_operationlog')
        )

class OperationLogViewSet(viewsets.ReadOnlyModelViewSet):
    """
    操作日志只读查询接口
    """
    queryset = OperationLog.objects.all()
    serializer_class = OperationLogSerializer
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_class = OperationLogFilter
    search_fields = ['username', 'module', 'action', 'path']
    ordering_fields = ['created_at', 'duration', 'response_code']
    ordering = ['-created_at']

    def get_permissions(self):
        """
        进行严格的权限检查：必须为管理员(is_staff)或拥有操作日志查看权限(operation_logs.view_operationlog / accounts.view_operationlog)
        """
        return [permissions.IsAuthenticated(), OperationLogReadPermission()]


class OperationLogCleanupAPIView(APIView):
    """手动立即清理一次过期操作日志。"""

    permission_classes = [permissions.IsAuthenticated, OperationLogCleanupPermission]

    def post(self, request):
        result = cleanup_operation_logs()
        return Response(result, status=status.HTTP_200_OK)


class OperationLogSettingAPIView(APIView):
    """操作日志自动清理设置。"""

    permission_classes = [permissions.IsAuthenticated, OperationLogReadPermission]

    def get(self, request):
        serializer = OperationLogSettingSerializer(OperationLogSetting.get_config())
        return Response(serializer.data, status=status.HTTP_200_OK)

    def patch(self, request):
        cleanup_permission = OperationLogCleanupPermission()
        if not cleanup_permission.has_permission(request, self):
            return Response({'detail': '无权修改操作日志清理设置'}, status=status.HTTP_403_FORBIDDEN)

        config = OperationLogSetting.get_config()
        serializer = OperationLogSettingSerializer(config, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)


class AnonymizationRulePermission(permissions.BasePermission):
    """脱敏规则管理权限：持有 knowledge.anonymize_document 权限的用户可操作"""
    def has_permission(self, request, view):
        if not request.user or request.user.is_anonymous:
            return False
        if request.user.is_superuser:
            return True
        return request.user.has_perm('knowledge.anonymize_document')


class AnonymizationRuleViewSet(viewsets.ModelViewSet):
    """脱敏规则 CRUD 接口"""

    queryset = AnonymizationRule.objects.all()
    serializer_class = AnonymizationRuleSerializer
    permission_classes = [permissions.IsAuthenticated, AnonymizationRulePermission]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name', 'entity_type', 'entity_label', 'description']
    ordering_fields = ['created_at', 'entity_type', 'score']
    ordering = ['-is_active', 'entity_type']

    @action(detail=False, methods=['post'], url_path='seed-defaults')
    def seed_defaults(self, request):
        """将内置默认规则写入数据库（已存在的同名规则也会更新中文标签）"""
        from utils.anonymization.chinese_recognizers import get_all_chinese_recognizers

        ENTITY_LABEL_MAP = {
            'PHONE_NUMBER': '手机号',
            'ID_CARD': '身份证号',
            'BANK_CARD': '银行卡号',
            'EMAIL_ADDRESS': '邮箱地址',
            'URL': '网址链接',
        }

        # 额外内置规则（不依赖 Presidio 识别器）
        EXTRA_RULES = [
            {
                'name': 'financial_account',
                'entity_type': 'FINANCIAL_ACCOUNT',
                'entity_label': '资金账户',
                'regex': r'\d{10,30}',
                'score': 0.7,
                'is_active': True,
                'description': '内置规则：匹配10-30位数字组成的资金账户号',
            },
        ]

        created_count = 0
        updated_count = 0
        for recognizer in get_all_chinese_recognizers():
            pattern = recognizer.patterns[0] if recognizer.patterns else None
            if not pattern:
                continue
            entity_type = recognizer.supported_entities[0] if recognizer.supported_entities else 'UNKNOWN'
            # 跳过已移除的类型
            if entity_type == 'LICENSE_PLATE':
                continue
            chinese_label = ENTITY_LABEL_MAP.get(entity_type, pattern.name)
            obj, created = AnonymizationRule.objects.get_or_create(
                name=pattern.name,
                defaults={
                    'entity_type': entity_type,
                    'entity_label': chinese_label,
                    'regex': pattern.regex,
                    'score': pattern.score,
                    'is_active': True,
                    'description': f'内置规则：{entity_type}',
                },
            )
            if created:
                created_count += 1
            else:
                if obj.entity_label != chinese_label:
                    obj.entity_label = chinese_label
                    obj.save(update_fields=['entity_label'])
                    updated_count += 1

        # 写入额外规则
        for rule_data in EXTRA_RULES:
            obj, created = AnonymizationRule.objects.get_or_create(
                name=rule_data['name'],
                defaults=rule_data,
            )
            if created:
                created_count += 1

        # 禁用车牌号规则
        AnonymizationRule.objects.filter(entity_type='LICENSE_PLATE').update(is_active=False)

        return Response(
            {'message': f'初始化完成：新增 {created_count} 条，更新 {updated_count} 条', 'created': created_count, 'updated': updated_count},
            status=status.HTTP_200_OK,
        )

    @action(detail=True, methods=['post'], url_path='test')
    def test_rule(self, request, pk=None):
        """测试单条规则对指定文本的匹配效果"""
        rule = self.get_object()
        text = request.data.get('text', '')
        if not text:
            return Response({'error': '请提供 text 字段'}, status=status.HTTP_400_BAD_REQUEST)
        import re
        try:
            matches = list(re.finditer(rule.regex, text))
        except re.error as e:
            return Response({'error': f'正则表达式错误: {e}'}, status=status.HTTP_400_BAD_REQUEST)
        results = [
            {'start': m.start(), 'end': m.end(), 'text': m.group()}
            for m in matches
        ]
        return Response({'matches': results, 'count': len(results)}, status=status.HTTP_200_OK)


class FileAnonymizeAPIView(APIView):
    """
    文件脱敏接口：上传文件 → 脱敏处理 → 下载脱敏后的文件。
    不写入数据库，仅做即时文件处理。
    支持 .txt, .md, .docx 格式。
    """
    permission_classes = [permissions.IsAuthenticated, AnonymizationRulePermission]
    parser_classes = [MultiPartParser]

    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    def post(self, request):
        uploaded = request.FILES.get('file')
        if not uploaded:
            return Response({'error': '请上传文件'}, status=status.HTTP_400_BAD_REQUEST)

        # 校验文件扩展名
        ext = os.path.splitext(uploaded.name)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return Response(
                {'error': f'不支持的文件格式 {ext}，仅支持 {", ".join(sorted(self.SUPPORTED_EXTENSIONS))}'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # 校验文件大小
        if uploaded.size > self.MAX_FILE_SIZE:
            return Response(
                {'error': f'文件过大，最大支持 {self.MAX_FILE_SIZE // 1024 // 1024}MB'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            from utils.anonymization.service import DocumentAnonymizer
            anonymizer = DocumentAnonymizer()

            if ext in ('.txt', '.md'):
                return self._process_text_file(uploaded, ext, anonymizer)
            elif ext == '.docx':
                return self._process_docx_file(uploaded, anonymizer)

        except Exception as e:
            logger.error("文件脱敏失败: %s", e, exc_info=True)
            return Response({'error': f'文件脱敏处理失败: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _process_text_file(self, uploaded, ext, anonymizer):
        """处理纯文本文件"""
        import chardet
        raw = uploaded.read()
        # 检测编码
        detected = chardet.detect(raw)
        encoding = detected.get('encoding') or 'utf-8'
        try:
            text = raw.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            text = raw.decode('utf-8', errors='replace')

        # 执行脱敏
        result = anonymizer.anonymize(text)

        # 构造响应
        anonymized_bytes = result.anonymized_text.encode('utf-8')
        base_name = os.path.splitext(uploaded.name)[0]
        output_name = f"{base_name}_anonymized{ext}"

        response = HttpResponse(anonymized_bytes, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{output_name}"'
        response['X-PII-Count'] = str(len(result.entities_found))
        return response

    def _process_docx_file(self, uploaded, anonymizer):
        """处理 Word 文档（.docx）"""
        from docx import Document
        from docx.oxml.ns import qn
        import copy

        doc = Document(io.BytesIO(uploaded.read()))
        total_pii_count = 0
        total_entities = []

        # 处理正文段落
        for para in doc.paragraphs:
            if para.text.strip():
                result = anonymizer.anonymize(para.text)
                total_pii_count += len(result.entities_found)
                total_entities.extend(result.entities_found)
                if result.entities_found:
                    # 保留段落格式，只替换文本
                    self._replace_paragraph_text(para, result.anonymized_text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            result = anonymizer.anonymize(para.text)
                            total_pii_count += len(result.entities_found)
                            total_entities.extend(result.entities_found)
                            if result.entities_found:
                                self._replace_paragraph_text(para, result.anonymized_text)

        # 输出脱敏后的 docx
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        base_name = os.path.splitext(uploaded.name)[0]
        output_name = f"{base_name}_anonymized.docx"

        response = HttpResponse(
            output_buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{output_name}"'
        response['X-PII-Count'] = str(total_pii_count)
        return response

    @staticmethod
    def _replace_paragraph_text(paragraph, new_text):
        """保留段落第一个 run 的格式，替换全部文本"""
        runs = paragraph.runs
        if not runs:
            return
        # 保留第一个 run 的格式，清除其他 run
        for run in runs[1:]:
            run.text = ''
        runs[0].text = new_text


class AnonymizedDocumentFilter(django_filters.FilterSet):
    status = django_filters.CharFilter(lookup_expr='exact')
    original_filename = django_filters.CharFilter(lookup_expr='icontains')

    class Meta:
        model = AnonymizedDocument
        fields = ['status']


class AnonymizationTemplateViewSet(viewsets.ModelViewSet):
    """脱敏模板 CRUD 接口"""

    queryset = AnonymizationTemplate.objects.all()
    serializer_class = AnonymizationTemplateSerializer
    permission_classes = [permissions.IsAuthenticated, AnonymizationRulePermission]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name', 'description']
    ordering_fields = ['created_at', 'updated_at', 'name']
    ordering = ['-updated_at']

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=True, methods=['post'], url_path='apply')
    def apply_template(self, request, pk=None):
        """将模板应用到指定文档"""
        template = self.get_object()
        document_id = request.data.get('document_id')
        if not document_id:
            return Response({'error': '请提供 document_id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            doc = AnonymizedDocument.objects.get(id=document_id)
        except AnonymizedDocument.DoesNotExist:
            return Response({'error': f'文档 {document_id} 不存在'}, status=status.HTTP_404_NOT_FOUND)
        doc.enabled_preset_types = list(template.enabled_preset_types)
        doc.custom_keywords = list(template.custom_keywords)
        doc.save(update_fields=['enabled_preset_types', 'custom_keywords'])
        serializer = AnonymizedDocumentSerializer(doc)
        return Response(serializer.data)

    @action(detail=False, methods=['post'], url_path='apply-multiple')
    def apply_multiple_templates(self, request):
        """将多个模板同时应用到指定文档（预设类型取并集，关键词按keyword去重合并）"""
        template_ids = request.data.get('template_ids', [])
        document_id = request.data.get('document_id')
        if not document_id:
            return Response({'error': '请提供 document_id'}, status=status.HTTP_400_BAD_REQUEST)
        if not template_ids or not isinstance(template_ids, list):
            return Response({'error': '请提供 template_ids 列表'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            doc = AnonymizedDocument.objects.get(id=document_id)
        except AnonymizedDocument.DoesNotExist:
            return Response({'error': f'文档 {document_id} 不存在'}, status=status.HTTP_404_NOT_FOUND)

        templates = AnonymizationTemplate.objects.filter(id__in=template_ids)
        if not templates.exists():
            return Response({'error': '未找到指定的模板'}, status=status.HTTP_404_NOT_FOUND)

        # 合并规则：预设类型取并集，关键词按keyword去重
        merged_preset_types = set()
        merged_keywords = {}  # keyword -> {keyword, replacement}
        for tpl in templates:
            for pt in (tpl.enabled_preset_types or []):
                merged_preset_types.add(pt)
            for kw in (tpl.custom_keywords or []):
                kw_key = kw.get('keyword', '') if isinstance(kw, dict) else kw
                if kw_key and kw_key not in merged_keywords:
                    merged_keywords[kw_key] = kw if isinstance(kw, dict) else {'keyword': kw, 'replacement': ''}

        doc.enabled_preset_types = sorted(merged_preset_types)
        doc.custom_keywords = list(merged_keywords.values())
        doc.save(update_fields=['enabled_preset_types', 'custom_keywords'])
        serializer = AnonymizedDocumentSerializer(doc)
        return Response(serializer.data)

    @action(detail=False, methods=['post'], url_path='seed-defaults')
    def seed_defaults(self, request):
        """初始化内置脱敏模板（不覆盖已有同名模板）"""
        DEFAULT_TEMPLATES = [
            {
                'name': '交易所专属标识全域消隐规则',
                'description': '将涉及特定机构的专有词汇替换为中性通用表述',
                'enabled_preset_types': [],
                'custom_keywords': [
                    {'keyword': '上海证券交易所', 'replacement': '相关证券交易所'},
                    {'keyword': '上交所', 'replacement': '相关交易所'},
                    {'keyword': '上证所', 'replacement': '相关交易所'},
                    {'keyword': '上证', 'replacement': '相关交易所'},
                    {'keyword': 'e投票', 'replacement': '线上投票'},
                    {'keyword': 'EP编码', 'replacement': '业务编码'},
                ],
            },
        ]

        created_count = 0
        for tpl_data in DEFAULT_TEMPLATES:
            obj, created = AnonymizationTemplate.objects.get_or_create(
                name=tpl_data['name'],
                defaults={
                    'description': tpl_data['description'],
                    'enabled_preset_types': tpl_data['enabled_preset_types'],
                    'custom_keywords': tpl_data['custom_keywords'],
                    'created_by': request.user,
                },
            )
            if created:
                created_count += 1
            else:
                # 更新已有模板的关键词和描述
                obj.description = tpl_data['description']
                obj.custom_keywords = tpl_data['custom_keywords']
                obj.save(update_fields=['description', 'custom_keywords'])
        return Response(
            {'message': f'初始化完成：新增 {created_count} 个模板，已更新已有模板', 'created': created_count},
            status=status.HTTP_200_OK,
        )


class AnonymizedDocumentViewSet(viewsets.ModelViewSet):
    """
    文档脱敏管理接口：上传、配置规则、执行脱敏、下载、删除。
    每个文档独立配置脱敏规则（预设PII类型 + 自定义关键词）。
    """
    queryset = AnonymizedDocument.objects.all()
    serializer_class = AnonymizedDocumentSerializer
    permission_classes = [permissions.IsAuthenticated, AnonymizationRulePermission]
    parser_classes = [MultiPartParser, JSONParser]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_class = AnonymizedDocumentFilter
    search_fields = ['original_filename']
    ordering_fields = ['created_at', 'status', 'file_size']
    ordering = ['-created_at']
    pagination_class = StandardPagination

    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    def create(self, request, *args, **kwargs):
        """上传文档（支持多文件）"""
        files = request.FILES.getlist('files')
        if not files:
            file = request.FILES.get('file')
            if file:
                files = [file]
        if not files:
            return Response({'error': '请上传至少一个文件'}, status=status.HTTP_400_BAD_REQUEST)

        created_docs = []
        errors = []

        for uploaded in files:
            ext = os.path.splitext(uploaded.name)[1].lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                errors.append({'file': uploaded.name, 'error': f'不支持的格式 {ext}'})
                continue
            if uploaded.size > self.MAX_FILE_SIZE:
                errors.append({'file': uploaded.name, 'error': '文件过大(最大10MB)'})
                continue

            doc = AnonymizedDocument.objects.create(
                original_file=uploaded,
                original_filename=uploaded.name,
                file_type=ext,
                file_size=uploaded.size,
                uploaded_by=request.user,
                # 默认启用所有预设类型
                enabled_preset_types=['PHONE_NUMBER', 'ID_CARD', 'EMAIL_ADDRESS', 'BANK_CARD', 'LICENSE_PLATE', 'URL'],
                custom_keywords=[],
            )
            created_docs.append(doc)

        serializer = AnonymizedDocumentSerializer(created_docs, many=True)
        return Response(
            {'docs': serializer.data, 'errors': errors, 'created_count': len(created_docs)},
            status=status.HTTP_201_CREATED,
        )

    def partial_update(self, request, *args, **kwargs):
        """更新文档的脱敏规则配置（仅允许修改 enabled_preset_types 和 custom_keywords）"""
        doc = self.get_object()
        allowed_fields = {'enabled_preset_types', 'custom_keywords'}
        data = {k: v for k, v in request.data.items() if k in allowed_fields}
        if not data:
            return Response({'error': '没有可更新的字段'}, status=status.HTTP_400_BAD_REQUEST)

        serializer = AnonymizedDocumentSerializer(doc, data=data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        """删除文档（同时删除原文件和脱敏文件）"""
        doc = self.get_object()
        # 删除物理文件
        if doc.original_file:
            try:
                doc.original_file.delete(save=False)
            except Exception:
                pass
        if doc.anonymized_file:
            try:
                doc.anonymized_file.delete(save=False)
            except Exception:
                pass
        doc.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['post'], url_path='execute')
    def execute_anonymization(self, request, pk=None):
        """执行脱敏：使用文档上配置的规则"""
        doc = self.get_object()

        # 支持多次脱敏：已脱敏文档可重复执行，每次从原始文件读取，报告覆盖最新
        # 检查规则配置
        if not doc.enabled_preset_types and not doc.custom_keywords:
            return Response({'error': '请先配置脱敏规则（勾选预设类型或添加关键词）'}, status=status.HTTP_400_BAD_REQUEST)

        doc.status = 'anonymizing'
        doc.save(update_fields=['status'])

        try:
            from utils.anonymization.service import DocumentAnonymizer
            from django.utils import timezone
            from django.core.files.base import ContentFile

            if doc.file_type == '.docx':
                # DOCX 必须按实际 XML 段落原位处理，不能先拍平再按行写回；
                # 否则页眉、页脚、目录、文本框会漏检，空段落也会导致内容错位。
                result, anonymized_bytes = self._anonymize_docx_content(
                    doc,
                    enabled_preset_types=doc.enabled_preset_types or [],
                    custom_keywords=doc.custom_keywords or [],
                )
                output_name = os.path.splitext(doc.original_filename)[0] + '_anonymized.docx'
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                text, _ = self._read_file_content(doc)
                result = DocumentAnonymizer.anonymize_with_config(
                    text=text,
                    enabled_preset_types=doc.enabled_preset_types or [],
                    custom_keywords=doc.custom_keywords or [],
                )
                anonymized_bytes = result.anonymized_text.encode('utf-8')
                output_name = os.path.splitext(doc.original_filename)[0] + '_anonymized' + doc.file_type
                content_type = 'text/plain'

            # 生成脱敏报告
            report = self._generate_report(result.entities_found)

            # 保存到模型
            doc.anonymized_file.save(output_name, ContentFile(anonymized_bytes), save=False)
            doc.anonymized_at = timezone.now()
            doc.anonymization_report = report
            doc.status = 'anonymized'
            doc.error_message = ''
            doc.save()

            serializer = AnonymizedDocumentSerializer(doc)
            return Response(serializer.data)

        except Exception as e:
            logger.error("文档脱敏失败 [doc_id=%s]: %s", doc.id, e, exc_info=True)
            doc.status = 'failed'
            doc.error_message = str(e)
            doc.save(update_fields=['status', 'error_message'])
            return Response({'error': f'脱敏失败: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'], url_path='download')
    def download_anonymized(self, request, pk=None):
        """下载脱敏后的文件"""
        doc = self.get_object()

        if not doc.anonymized_file:
            logger.warning("下载失败 [doc_id=%s]: 脱敏文件不存在", doc.id)
            return Response({'error': '文档尚未脱敏，无法下载'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            doc.anonymized_file.seek(0)
            file_content = doc.anonymized_file.read()
        except Exception as e:
            logger.error("读取脱敏文件失败 [doc_id=%s, path=%s]: %s", doc.id, doc.anonymized_file.name, e, exc_info=True)
            return Response(
                {'error': f'读取脱敏文件失败: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        if not file_content:
            logger.warning("下载失败 [doc_id=%s, path=%s]: 文件内容为空", doc.id, doc.anonymized_file.name)
            return Response({'error': '脱敏文件内容为空，请重新执行脱敏'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        from django.utils import timezone
        date_str = timezone.now().strftime('%Y%m%d')
        orig_name, orig_ext = os.path.splitext(doc.original_filename)
        output_name = f'已脱敏_{orig_name}_{date_str}{orig_ext}'
        logger.info("下载脱敏文件 [doc_id=%s, filename=%s, size=%d bytes]", doc.id, output_name, len(file_content))

        response = HttpResponse(file_content, content_type='application/octet-stream')
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{urllib.parse.quote(output_name)}"
        response['Content-Length'] = len(file_content)
        return response

    @action(detail=True, methods=['post'], url_path='reset')
    def reset_anonymization(self, request, pk=None):
        """重置脱敏状态（删除脱敏结果，回到待脱敏）"""
        doc = self.get_object()
        if doc.anonymized_file:
            try:
                doc.anonymized_file.delete(save=False)
            except Exception:
                pass
        doc.status = 'pending'
        doc.anonymized_file = None
        doc.anonymized_at = None
        doc.anonymization_report = None
        doc.error_message = ''
        doc.save()
        serializer = AnonymizedDocumentSerializer(doc)
        return Response(serializer.data)

    # ----- 内部方法 -----

    @staticmethod
    def _iter_docx_paragraphs(document):
        """遍历正文、表格、目录、文本框以及所有页眉页脚中的段落。"""
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        roots = [document.element.body]
        for section in document.sections:
            roots.extend([
                section.header._element,
                section.first_page_header._element,
                section.even_page_header._element,
                section.footer._element,
                section.first_page_footer._element,
                section.even_page_footer._element,
            ])

        seen_roots = set()
        seen_paragraphs = set()
        for root in roots:
            if root in seen_roots:
                continue
            seen_roots.add(root)
            for element in root.iter(qn('w:p')):
                if element in seen_paragraphs:
                    continue
                seen_paragraphs.add(element)
                yield Paragraph(element, document)

    @staticmethod
    def _paragraph_text_nodes(paragraph):
        """获取当前段落直属的文本节点，排除文本框外层段落包含的内层段落。"""
        from docx.oxml.ns import qn

        paragraph_tag = qn('w:p')
        nodes = []
        for node in paragraph._p.iter(qn('w:t')):
            ancestor = node.getparent()
            while ancestor is not None and ancestor.tag != paragraph_tag:
                ancestor = ancestor.getparent()
            if ancestor is paragraph._p:
                nodes.append(node)
        return nodes

    @classmethod
    def _anonymize_docx_content(cls, doc, enabled_preset_types, custom_keywords):
        """按 DOCX 节点原位脱敏，返回汇总结果和输出文件字节。"""
        from docx import Document as DocxDocument
        from utils.anonymization.service import DocumentAnonymizer

        doc.original_file.seek(0)
        document = DocxDocument(io.BytesIO(doc.original_file.read()))
        paragraph_entries = []
        for paragraph in cls._iter_docx_paragraphs(document):
            text_nodes = cls._paragraph_text_nodes(paragraph)
            paragraph_text = ''.join(node.text or '' for node in text_nodes)
            if paragraph_text.strip() and text_nodes:
                paragraph_entries.append((text_nodes, paragraph_text))

        # DOCX XML 不允许出现这些控制字符，因此它们不可能来自原文，也不会被
        # 手机号、邮箱、URL 等文本规则命中。此前边界中包含随机 UUID；当 UUID
        # 恰好含有 11 位手机号样式的数字时，边界会被脱敏并导致段落无法拆回。
        boundary = '\x00\x01\x02\x03'
        combined_text = boundary.join(text for _, text in paragraph_entries)
        result = DocumentAnonymizer.anonymize_with_config(
            text=combined_text,
            enabled_preset_types=enabled_preset_types,
            custom_keywords=custom_keywords,
        )
        anonymized_paragraphs = result.anonymized_text.split(boundary)
        if len(anonymized_paragraphs) != len(paragraph_entries):
            raise ValueError('DOCX 脱敏后的段落边界异常，已停止写入以避免文档内容错位。')

        for (text_nodes, original_text), anonymized_text in zip(
            paragraph_entries, anonymized_paragraphs
        ):
            if anonymized_text == original_text:
                continue
            text_nodes[0].text = anonymized_text
            for node in text_nodes[1:]:
                node.text = ''

        output_buffer = io.BytesIO()
        document.save(output_buffer)
        return result, output_buffer.getvalue()

    @staticmethod
    def _read_file_content(doc):
        """读取文档内容，返回 (text, is_docx)"""
        doc.original_file.seek(0)
        if doc.file_type == '.docx':
            from docx import Document as DocxDocument
            try:
                file_bytes = doc.original_file.read()
                d = DocxDocument(io.BytesIO(file_bytes))
            except Exception as e:
                raise ValueError(
                    f'无法解析 .docx 文件（可能文件已损坏或格式不兼容）: {e}'
                ) from e
            paragraphs = [p.text for p in d.paragraphs]
            # 表格内容也提取
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                paragraphs.append(p.text)
            return '\n'.join(paragraphs), True
        else:
            raw = doc.original_file.read()
            try:
                import chardet
                detected = chardet.detect(raw)
                encoding = detected.get('encoding') or 'utf-8'
                text = raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                text = raw.decode('utf-8', errors='replace')
            return text, False

    @staticmethod
    def _write_docx_content(doc, anonymized_text):
        """将脱敏后文本写回 docx 格式，返回 bytes"""
        from docx import Document as DocxDocument

        doc.original_file.seek(0)
        d = DocxDocument(io.BytesIO(doc.original_file.read()))
        lines = anonymized_text.split('\n')
        line_idx = 0

        # 替换正文段落
        for para in d.paragraphs:
            if para.text.strip() and line_idx < len(lines):
                runs = para.runs
                if runs:
                    for run in runs[1:]:
                        run.text = ''
                    runs[0].text = lines[line_idx]
                line_idx += 1

        # 替换表格
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip() and line_idx < len(lines):
                            runs = para.runs
                            if runs:
                                for run in runs[1:]:
                                    run.text = ''
                                runs[0].text = lines[line_idx]
                            line_idx += 1

        output_buffer = io.BytesIO()
        d.save(output_buffer)
        return output_buffer.getvalue()

    @staticmethod
    def _generate_report(entities_found):
        """生成脱敏报告"""
        # 按类型分组统计
        type_counts = {}
        for entity in entities_found:
            et = entity['entity_type']
            label = entity.get('entity_label', et)
            if et not in type_counts:
                type_counts[et] = {'entity_type': et, 'entity_label': label, 'count': 0, 'examples': []}
            type_counts[et]['count'] += 1
            if len(type_counts[et]['examples']) < 5:
                type_counts[et]['examples'].append(entity['text_snippet'])

        return {
            'total_count': len(entities_found),
            'details': list(type_counts.values()),
            'all_entities': entities_found[:100],  # 最多保留 100 条明细
        }
