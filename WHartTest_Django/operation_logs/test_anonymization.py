import io
from types import SimpleNamespace

from django.core.files.base import ContentFile
from django.test import SimpleTestCase
from docx import Document
from docx.oxml import OxmlElement

from operation_logs.views import AnonymizedDocumentViewSet
from utils.anonymization.service import DocumentAnonymizer


class CustomKeywordAnonymizationTests(SimpleTestCase):
    def test_literal_keyword_matching_is_case_insensitive_and_not_deleted(self):
        result = DocumentAnonymizer.anonymize_with_config(
            text='Acme ACME acme Acme.*',
            enabled_preset_types=[],
            custom_keywords=[{'keyword': 'acme', 'replacement': ''}],
        )

        self.assertEqual(result.anonymized_text, '<关键词> <关键词> <关键词> <关键词>.*')
        self.assertEqual(len(result.entities_found), 4)

    def test_custom_replacement_is_used(self):
        result = DocumentAnonymizer.anonymize_with_config(
            text='项目代号 SECRET，复核 secret。',
            enabled_preset_types=[],
            custom_keywords=[{'keyword': 'Secret', 'replacement': '[已脱敏]'}],
        )

        self.assertEqual(result.anonymized_text, '项目代号 [已脱敏]，复核 [已脱敏]。')


class DocxAnonymizationScopeTests(SimpleTestCase):
    @staticmethod
    def _append_textbox(paragraph, text):
        textbox = OxmlElement('w:txbxContent')
        inner_paragraph = OxmlElement('w:p')
        run = OxmlElement('w:r')
        text_node = OxmlElement('w:t')
        text_node.text = text
        run.append(text_node)
        inner_paragraph.append(run)
        textbox.append(inner_paragraph)
        paragraph.add_run()._r.append(textbox)

    def test_docx_body_table_toc_textbox_header_and_footer_are_processed(self):
        document = Document()
        document.add_paragraph('正文 Secret')
        document.add_paragraph('目录 SECRET')
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = '表格 secret'
        self._append_textbox(document.add_paragraph(), '文本框 SeCrEt')
        document.sections[0].header.paragraphs[0].text = '页眉 SECRET'
        document.sections[0].footer.paragraphs[0].text = '页脚 secret'

        source = io.BytesIO()
        document.save(source)
        stored_document = SimpleNamespace(original_file=ContentFile(source.getvalue()))

        result, output = AnonymizedDocumentViewSet._anonymize_docx_content(
            stored_document,
            enabled_preset_types=[],
            custom_keywords=[{'keyword': 'secret', 'replacement': '[敏感]'}],
        )

        processed = Document(io.BytesIO(output))
        all_text = []
        for paragraph in AnonymizedDocumentViewSet._iter_docx_paragraphs(processed):
            nodes = AnonymizedDocumentViewSet._paragraph_text_nodes(paragraph)
            all_text.append(''.join(node.text or '' for node in nodes))
        combined = '\n'.join(all_text)

        self.assertNotIn('secret', combined.casefold())
        self.assertGreaterEqual(combined.count('[敏感]'), 6)
        self.assertEqual(len(result.entities_found), 6)

    def test_phone_number_rule_cannot_modify_docx_paragraph_boundary(self):
        document = Document()
        document.add_paragraph('联系人手机号 13800138000')
        document.add_paragraph('备用手机号 13900139000')

        source = io.BytesIO()
        document.save(source)
        stored_document = SimpleNamespace(original_file=ContentFile(source.getvalue()))

        result, output = AnonymizedDocumentViewSet._anonymize_docx_content(
            stored_document,
            enabled_preset_types=['PHONE_NUMBER'],
            custom_keywords=[],
        )

        processed = Document(io.BytesIO(output))
        combined = '\n'.join(paragraph.text for paragraph in processed.paragraphs)

        self.assertNotIn('13800138000', combined)
        self.assertNotIn('13900139000', combined)
        self.assertEqual(combined.count('<PHONE_NUMBER>'), 2)
        self.assertEqual(len(result.entities_found), 2)
